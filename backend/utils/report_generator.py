import json
import logging
import os
import zipfile

from docx import Document
from docx.shared import Inches
from flask import url_for

from services.server_services import open_file
from utils.utils import BASE_DIR, REPORTS_DIR, TEMP_DIR, TEMPLATES_DIR, get_path, parse_log_date

def validate_cluster_names(cluster_names):
    """Validates the cluster names input."""
    if not cluster_names:
        return False, "No cluster names provided."
    if not all(isinstance(name, str) and name.strip() for name in cluster_names):
        return False, "All cluster names must be non-empty strings."
    return True, None


def replace_placeholders(doc, placeholders):
    """
    Replaces text placeholders in both paragraphs and tables of the Word document.
    """
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)


def insert_images(doc, image_placeholders, image_folder):
    """Inserts images in the Word document replacing placeholders."""
    for paragraph in doc.paragraphs:
        for placeholder, image_name in image_placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, "")
                image_path = os.path.join(image_folder, image_name)
                if os.path.exists(image_path):
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(5))


def update_summary_placeholders(summary_placeholders, summary_warning_counts, cluster_index, placeholders):
    summary_placeholders[f"[cluster_{cluster_index}]"] = placeholders.get("[cluster_name]", "Unknown")
    summary_placeholders[f"[checklist_1_cluster_{cluster_index}]"] = "✓" if placeholders.get("[warning_1]") == "OK" else "⚠"
    summary_placeholders[f"[checklist_2_cluster_{cluster_index}]"] = "✓" if placeholders.get("[warning_2]") == "OK" else "⚠"
    summary_warning_counts[f"[state_sum_cluster_{cluster_index}]"] = sum(
        1 for key in ["[message_1]", "[message_2]"] if placeholders.get(key)
    )


def generate_summary_report(summary_template_path, collection_name, summary_placeholders, summary_warning_counts):
    summary_doc = Document(summary_template_path)
    replace_placeholders(summary_doc, summary_placeholders)
    for key, value in summary_warning_counts.items():
        replace_placeholders(summary_doc, {key: str(value)})

    summary_output_path = get_path(REPORTS_DIR, collection_name, 'Summary_Report.docx')
    summary_doc.save(summary_output_path)
   

def process_cluster_report(template_path, collection_name, cluster_name, data):
    """
    Process the report for a single cluster.
    """
    doc = Document(template_path)

    # Update placeholders and document content
    placeholders = generate_placeholders(collection_name, cluster_name, data)
    replace_placeholders(doc, placeholders)

    # Insert logs, metadata, and images
    insert_logs(doc, data, collection_name, cluster_name)
    insert_logs_metadata(doc, data, collection_name, cluster_name)
    insert_images(
        doc,
        generate_image_placeholders(data),
        get_path(BASE_DIR, collection_name, cluster_name, "charts"),
    )

    # Save the document
    output_path = get_path(REPORTS_DIR, collection_name, f"{cluster_name}_Report.docx")
    doc.save(output_path)

    return output_path, placeholders


def generate_placeholders(collection_name, cluster_name, data):
    placeholders = {
        "[cluster_name]": data.get("cluster_name", ""),
        **{
            f"[server_{j+1}_name]": server.get("server_name", "")
            for j, server in enumerate(data.get("servers", []))
        },
    }

    # Add warnings and messages
    warning_result = calculate_warnings(
        collection_name, cluster_name, "2_req-resp.json", error_file="6_error-count.json"
    )
    placeholders.update(warning_result)
    return placeholders


def calculate_warnings(collection_name, cluster_name, file_name, error_file=None):
    try:
        request_data = load_json_summary(collection_name, cluster_name, file_name)
        if not request_data:
            return default_warnings("Failed to retrieve request/response data.")

        request_count, response_count = get_request_response_counts(request_data)
        total_count = request_count + response_count

        if total_count == 0:
            return default_warnings("No Request or Response data available.")

        warning_1, message_1 = evaluate_percentages(request_count, response_count, total_count)
        warning_2, message_2 = evaluate_errors(collection_name, cluster_name, error_file)

        return {
            "[warning_1]": warning_1,
            "[message_1]": message_1,
            "[warning_2]": warning_2,
            "[message_2]": message_2,
        }
    except Exception as e:
        return default_warnings(f"Unexpected error: {e}")


def load_json(collection_name, cluster_name, folder_name, file_name):
    response, status_code = open_file(collection_name, cluster_name, file_name, "server", folder_name)
    return response.get_json() if status_code == 200 else None

def load_json_summary(collection_name, cluster_name, file_name):
    response, status_code = open_file(collection_name, cluster_name, file_name, "summary")
    return response.get_json() if status_code == 200 else None


def default_warnings(message):
    return {
        "[warning_1]": "Warning",
        "[message_1]": message,
        "[warning_2]": "Warning",
        "[message_2]": "",
    }


def get_request_response_counts(data):
    request_count = next((item["Count"] for item in data if item["Operation"] == "request"), 0)
    response_count = next((item["Count"] for item in data if item["Operation"] == "response"), 0)
    return request_count, response_count


def evaluate_percentages(request_count, response_count, total_count):
    request_percentage = (request_count / total_count) * 100
    response_percentage = (response_count / total_count) * 100
    if abs(response_percentage - request_percentage) > 10:
        return "Warning", "The percentage difference between Request and Response exceeds 10%."
    return "OK", ""


def evaluate_errors(collection_name, cluster_name, error_file):
    if not error_file:
        return "OK", ""
    error_data = load_json_summary(collection_name, cluster_name, error_file)
    if not error_data:
        return "Warning", f"Failed to retrieve data from {error_file}."
    total_errors = sum(item["Count"] for item in error_data)
    if total_errors > 10:
        most_frequent_error = max(error_data, key=lambda x: x["Count"])["Errors"]
        return "Warning", f"Most frequent error: '{most_frequent_error}'."
    return "OK", ""


def insert_logs(doc, data, collection_name, cluster_name):
    servers = data.get("servers", [])
    for j, server in enumerate(servers):
        log_content = process_log_file(collection_name, cluster_name, server.get("server_name"), "0_listing-audit-logs.json")
        replace_placeholders(doc, {f"[log_{j+1}]": log_content})

def process_log_file(collection_name, cluster_name, server_name, log_file):
    log_file_path = get_path(BASE_DIR, collection_name, cluster_name, "servers", server_name, log_file)
    output_file_path = get_path(TEMPLATES_DIR, "log-content.txt")
    convert_json_to_text(log_file_path, output_file_path)
    return read_file(output_file_path)


def convert_json_to_text(json_file_path, output_file_path):
    with open(json_file_path, "r") as json_file:
        data = json.load(json_file)
    with open(output_file_path, "w") as output_file:
        output_file.write(
            f"total {data['total']}\n"
            + "\n".join(
                f"{log['permissions']} {log['links']} {log['owner']} {log['group']} {log['size']:>8} {log['date']} {log['name']}"
                for log in data["logs"]
            )
        )


def read_file(file_path):
    with open(file_path, "r") as file:
        return file.read()


def insert_logs_metadata(doc, data, collection_name,  cluster_name):
    servers = data.get("servers", [])
    for j, server in enumerate(servers):
        metadata = extract_log_metadata(collection_name, cluster_name, server.get("server_name"), "0_listing-audit-logs.json", j)
        if metadata:
            replace_placeholders(doc, metadata)


def extract_log_metadata(collection_name, cluster_name, server_name, log_file, index_server):
    j = index_server + 1
    log_data = load_json(collection_name, cluster_name, server_name, log_file)
    if not log_data:
        return {}
    logs = log_data.get("logs", [])
    dates = [parse_log_date(log.get("date")) for log in logs if log.get("date")]
    return {
        f"[ps_{j}]": min(dates).strftime("%Y-%m-%d") if dates else "",
        f"[pe_{j}]": max(dates).strftime("%Y-%m-%d") if dates else "",
        f"[nof_{j}]": str(len(log_data.get("logs", []))),
    }


def generate_image_placeholders(data):
    return {
        f"[chart_{j+1}]": chart
        for j, chart in enumerate(data.get("charts", []))
    }


def wrap_report(collection_name):    
    try:
        # Name of the ZIP file based on the collection name
        zip_filename = f"{collection_name}.zip"
        zip_path = get_path(TEMP_DIR, zip_filename)
        
        # Path to the collection folder
        collection_path = get_path(REPORTS_DIR, collection_name)

        # Validate if the collection folder exists
        if not os.path.isdir(collection_path):
            logging.error(f"Collection '{collection_name}' does not exist in {REPORTS_DIR}.")
            raise FileNotFoundError(f"Collection '{collection_name}' does not exist in {REPORTS_DIR}.")

        # Create ZIP file
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, dirs, files in os.walk(collection_path):
                for file in files:
                    file_path = os.path.join(root, file)  # Full path to the file
                    arcname = os.path.relpath(file_path, collection_path)  # Relative path inside the ZIP
                    zipf.write(file_path, arcname)

        # Generate download URL
        download_url = url_for('download_zip', filename=zip_filename, _external=True)
        return download_url

    except Exception as e:
        logging.error(f"Error while creating ZIP file: {str(e)}")
        raise
